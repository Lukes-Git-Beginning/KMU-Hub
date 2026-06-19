"""
Generates the Gesellschaftervertrag for Zentria UG as a formatted Word document.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(11)
style_normal.paragraph_format.space_after = Pt(6)
style_normal.paragraph_format.line_spacing = 1.15

style_h1 = doc.styles['Heading 1']
style_h1.font.name = 'Times New Roman'
style_h1.font.size = Pt(22)
style_h1.font.bold = True
style_h1.font.color.rgb = RGBColor(0, 0, 0)
style_h1.paragraph_format.space_before = Pt(0)
style_h1.paragraph_format.space_after = Pt(6)

style_h2 = doc.styles['Heading 2']
style_h2.font.name = 'Times New Roman'
style_h2.font.size = Pt(13)
style_h2.font.bold = True
style_h2.font.color.rgb = RGBColor(0, 0, 0)
style_h2.paragraph_format.space_before = Pt(18)
style_h2.paragraph_format.space_after = Pt(6)

style_h3 = doc.styles['Heading 3']
style_h3.font.name = 'Times New Roman'
style_h3.font.size = Pt(11)
style_h3.font.bold = True
style_h3.font.italic = True
style_h3.font.color.rgb = RGBColor(0, 0, 0)
style_h3.paragraph_format.space_before = Pt(12)
style_h3.paragraph_format.space_after = Pt(4)


def add_paragraph(text, bold=False, italic=False, indent=False, align=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p


def add_numbered(number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(f"({number})  {text}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.8)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(f"–  {text}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        # Gray background
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'E8E8E8',
            qn('w:val'): 'clear'
        })
        shading.append(shading_elem)
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '0',
            qn('w:color'): '999999'
        })
        borders.append(elem)
    tblPr.append(borders)
    doc.add_paragraph()  # spacing
    return table


# ═══════════════════════════════════════════════
#  DOCUMENT CONTENT
# ═══════════════════════════════════════════════

# Title page
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GESELLSCHAFTERVERTRAG")
run.font.name = 'Times New Roman'
run.font.size = Pt(28)
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Zentria UG (haftungsbeschränkt)")
run.font.name = 'Times New Roman'
run.font.size = Pt(18)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— ENTWURF —")
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.italic = True
run.font.color.rgb = RGBColor(180, 0, 0)

doc.add_paragraph()
doc.add_paragraph()

add_paragraph("Dieser Entwurf ist zur Vorlage beim Notar bestimmt.", italic=True, align='center')
add_paragraph("Die notarielle Beurkundung ist für die Rechtswirksamkeit zwingend erforderlich.", italic=True, align='center')

doc.add_paragraph()
doc.add_paragraph()

add_paragraph("Erstellt: März 2026", align='center')
add_paragraph("Gesellschafter: Luke [Nachname], Darien [Nachname], Nico [Nachname]", align='center')

doc.add_page_break()

# ── INHALTSVERZEICHNIS ──
doc.add_heading("Inhaltsverzeichnis", level=1)
doc.add_paragraph()

toc_items = [
    "§ 1    Firma und Sitz",
    "§ 2    Gegenstand des Unternehmens",
    "§ 3    Stammkapital und Geschäftsanteile",
    "§ 4    Geschäftsführung und Vertretung",
    "§ 5    Gesellschafterversammlungen",
    "§ 6    Beschlussfassung",
    "§ 7    Vesting der Geschäftsanteile",
    "§ 8    Ausscheiden eines Gesellschafters (Good Leaver / Bad Leaver)",
    "§ 9    Verfügung über Geschäftsanteile",
    "§ 10  Geistiges Eigentum (Intellectual Property)",
    "§ 11  Geheimhaltung",
    "§ 12  Wettbewerbsverbot",
    "§ 13  Gewinnverwendung und Rücklagen",
    "§ 14  Vergütung der Geschäftsführer",
    "§ 15  Deadlock-Regelung",
    "§ 16  Kündigung und Austritt",
    "§ 17  Umwandlung in GmbH",
    "§ 18  Auflösung und Liquidation",
    "§ 19  Schlussbestimmungen",
    "",
    "Anlage 1 — Vorbestehendes geistiges Eigentum",
    "Anlage 2 — Gesellschafterliste",
]
for item in toc_items:
    if item == "":
        doc.add_paragraph()
        continue
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════
# § 1
# ════════════════════════════════════════
doc.add_heading("§ 1 — Firma und Sitz", level=2)

add_numbered(1, "Die Firma der Gesellschaft lautet: Zentria UG (haftungsbeschränkt).")
add_numbered(2, "Sitz der Gesellschaft ist _______________________ [Stadt eintragen].")
add_numbered(3, "Das Geschäftsjahr ist das Kalenderjahr. Das erste Geschäftsjahr beginnt mit der Eintragung ins Handelsregister und endet am 31. Dezember desselben Jahres.")

# ════════════════════════════════════════
# § 2
# ════════════════════════════════════════
doc.add_heading("§ 2 — Gegenstand des Unternehmens", level=2)

add_numbered(1, "Gegenstand des Unternehmens ist:")
add_bullet("die Entwicklung, der Vertrieb und der Betrieb von Software, insbesondere cloudbasierter Unternehmenssoftware (SaaS) sowie selbst gehosteter Lösungen für kleine und mittlere Unternehmen")
add_bullet("die Erbringung von IT-Dienstleistungen, Beratung und Support")
add_bullet("der Handel mit Software und digitalen Produkten")

add_numbered(2, "Die Gesellschaft ist zu allen Geschäften und Maßnahmen berechtigt, die dem Gegenstand des Unternehmens dienen oder ihn fördern.")

# ════════════════════════════════════════
# § 3
# ════════════════════════════════════════
doc.add_heading("§ 3 — Stammkapital und Geschäftsanteile", level=2)

add_numbered(1, "Das Stammkapital der Gesellschaft beträgt 300,00 EUR (dreihundert Euro).")
add_numbered(2, "Das Stammkapital wird wie folgt übernommen:")

add_table(
    ["Nr.", "Gesellschafter", "Geschäftsanteil", "Nennbetrag", "Einlage"],
    [
        ["1", "Luke [Nachname]", "33,3 %", "100,00 EUR", "100,00 EUR"],
        ["2", "Darien [Nachname]", "33,3 %", "100,00 EUR", "100,00 EUR"],
        ["3", "Nico [Nachname]", "33,3 %", "100,00 EUR", "100,00 EUR"],
    ]
)

add_numbered(3, "Die Einlagen sind in voller Höhe sofort in bar zu leisten.")
add_numbered(4, "Die Bildung von Geschäftsanteilen mit unterschiedlichen Nennbeträgen ist zulässig.")

# ════════════════════════════════════════
# § 4
# ════════════════════════════════════════
doc.add_heading("§ 4 — Geschäftsführung und Vertretung", level=2)

add_numbered(1, "Die Gesellschaft hat drei Geschäftsführer:")
add_bullet("Luke [Nachname]")
add_bullet("Darien [Nachname]")
add_bullet("Nico [Nachname]")

add_numbered(2, "Vertretungsregelung: Die Gesellschaft wird durch zwei Geschäftsführer gemeinschaftlich vertreten (Gesamtvertretung). Kein Geschäftsführer ist einzelvertretungsberechtigt.")

add_numbered(3, "Alleinentscheidung im Tagesgeschäft: Für Geschäfte des laufenden Betriebs mit einem Wert von bis zu 500,00 EUR ist jeder Geschäftsführer einzeln handlungsbefugt. Dies umfasst insbesondere:")
add_bullet("Anschaffungen für den laufenden Betrieb (Software-Lizenzen, Hosting, Bürobedarf)")
add_bullet("Kleinere Reparaturen und Wartungen")

add_numbered(4, "Zustimmung der Gesellschafterversammlung ist erforderlich für:")
add_bullet("Ankauf, Verkauf und Belastung von Grundstücken")
add_bullet("Abschluss von Mietverträgen, Dienstverträgen und Darlehensverträgen jeglicher Art")
add_bullet("Übernahme von Garantien, Haftungen und Bürgschaften")
add_bullet("Pensionszusagen, soweit die Gesellschaft nicht schon durch Betriebsvereinbarung hierzu verpflichtet ist")
add_bullet("Aufnahme neuer Gesellschafter oder Abtretung von Geschäftsanteilen")
add_bullet("Einstellung und Kündigung von Mitarbeitern")
add_bullet("Erteilung und Widerruf von Prokura")
add_bullet("Gründung von oder Beteiligung an anderen Unternehmen")
add_bullet("Investitionen und Einzelausgaben über 2.000,00 EUR")
add_bullet("Aufnahme neuer Geschäftsfelder oder wesentliche Änderung des Unternehmensgegenstands")
add_bullet("Abschluss, Änderung oder Kündigung von Verträgen mit Gesellschaftern oder nahestehenden Personen")
add_bullet("Festlegung oder Änderung der Geschäftsführer-Vergütung")
add_bullet("Verkauf oder Lizenzierung von wesentlichem geistigem Eigentum (Software, Markenrechte, Patente)")
add_bullet("Umwandlung der Gesellschaft (z. B. UG in GmbH)")

add_numbered(5, "Jeder Geschäftsführer ist von den Beschränkungen des § 181 BGB (Insichgeschäft) nicht befreit.")

add_numbered(6, "Die Geschäftsführer sind verpflichtet, die Geschäfte der Gesellschaft mit der Sorgfalt eines ordentlichen Kaufmanns zu führen und sich gegenseitig über alle wesentlichen Geschäftsvorfälle zu informieren.")

# ════════════════════════════════════════
# § 5
# ════════════════════════════════════════
doc.add_heading("§ 5 — Gesellschafterversammlungen", level=2)

add_numbered(1, "Ordentliche Gesellschafterversammlung: Mindestens einmal jährlich innerhalb der ersten acht Monate des Geschäftsjahres (Jahresversammlung).")

add_numbered(2, "Außerordentliche Gesellschafterversammlungen sind einzuberufen, wenn:")
add_bullet("ein Geschäftsführer dies für erforderlich hält")
add_bullet("ein Gesellschafter dies unter Angabe von Gründen verlangt")
add_bullet("es das Wohl der Gesellschaft erfordert")

add_numbered(3, "Die Einberufung erfolgt durch einen Geschäftsführer mit einer Frist von 14 Tagen unter Angabe der Tagesordnung. Die Einberufung kann per E-Mail an die zuletzt bekannte E-Mail-Adresse erfolgen.")

add_numbered(4, "Quartalstreffen: Die Gesellschafter treffen sich zusätzlich einmal pro Quartal zur Besprechung der Geschäftsentwicklung, Strategie und Rollenverteilung. Diese Treffen sind formlos, aber zu protokollieren.")

# ════════════════════════════════════════
# § 6
# ════════════════════════════════════════
doc.add_heading("§ 6 — Beschlussfassung", level=2)

add_numbered(1, "Die Gesellschafterversammlung ist beschlussfähig, wenn alle Gesellschafter anwesend oder vertreten sind. Bei Abwesenheit eines Gesellschafters ist eine neue Versammlung mit gleicher Tagesordnung binnen 14 Tagen einzuberufen, die unabhängig von der Anwesenheit beschlussfähig ist.")

add_numbered(2, "Beschlüsse werden mit folgenden Mehrheiten gefasst:")

add_table(
    ["Beschlussgegenstand", "Erforderliche Mehrheit"],
    [
        ["Tagesgeschäft, laufende Entscheidungen", "Einfache Mehrheit (2 von 3)"],
        ["Feststellung des Jahresabschlusses", "Einfache Mehrheit (2 von 3)"],
        ["Gewinnverwendung", "Einfache Mehrheit (2 von 3)"],
        ["Bestellung / Abberufung von Geschäftsführern", "75 % der Stimmen"],
        ["Aufnahme neuer Gesellschafter", "Einstimmig (100 %)"],
        ["Änderung des Gesellschaftervertrags", "Einstimmig (100 %)"],
        ["Verkauf des Unternehmens (ganz oder wesentliche Teile)", "Einstimmig (100 %)"],
        ["Auflösung der Gesellschaft", "Einstimmig (100 %)"],
        ["Kapitalerhöhung / Ausgabe neuer Anteile", "Einstimmig (100 %)"],
    ]
)

add_numbered(3, "Jeder Euro des Nennbetrags eines Geschäftsanteils gewährt eine Stimme.")
add_numbered(4, "Beschlussfassung im Umlaufverfahren (schriftlich, per E-Mail) ist zulässig, wenn alle Gesellschafter diesem Verfahren zustimmen.")

# ════════════════════════════════════════
# § 7
# ════════════════════════════════════════
doc.add_heading("§ 7 — Vesting der Geschäftsanteile", level=2)

add_numbered(1, "Zweck: Die Geschäftsanteile unterliegen einer zeitlichen Bindung (Vesting), um sicherzustellen, dass alle Gesellschafter langfristig zum Erfolg der Gesellschaft beitragen.")
add_numbered(2, "Vesting-Zeitraum: 48 Monate (4 Jahre) ab Eintragung der Gesellschaft ins Handelsregister.")
add_numbered(3, "Cliff: Während der ersten 12 Monate (Cliff-Periode) werden keine Anteile gevestet. Scheidet ein Gesellschafter vor Ablauf der Cliff-Periode aus, fallen sämtliche Anteile an die Gesellschaft zurück (Einziehung).")
add_numbered(4, "Vesting nach dem Cliff: Nach Ablauf der Cliff-Periode vesten die Anteile linear monatlich über die verbleibenden 36 Monate. Pro Monat vestet 1/36 der Gesamtanteile des jeweiligen Gesellschafters.")

add_numbered(5, "Vesting-Tabelle (pro Gesellschafter bei 33,3 %):")

add_table(
    ["Zeitpunkt", "Gevestet", "Ungevestet"],
    [
        ["Monat 0–12 (Cliff)", "0 %", "33,3 %"],
        ["Monat 13", "0,93 %", "32,37 %"],
        ["Monat 24", "11,1 %", "22,2 %"],
        ["Monat 36", "22,2 %", "11,1 %"],
        ["Monat 48", "33,3 %", "0 %"],
    ]
)

add_numbered(6, "Beschleunigtes Vesting (Acceleration): Sämtliche Anteile eines Gesellschafters vesten sofort und vollständig bei:")
add_bullet("Verkauf der Gesellschaft (Change of Control) mit Zustimmung aller Gesellschafter")
add_bullet("Tod oder dauerhafter Erwerbsunfähigkeit des Gesellschafters")

# ════════════════════════════════════════
# § 8
# ════════════════════════════════════════
doc.add_heading("§ 8 — Ausscheiden eines Gesellschafters", level=2)
doc.add_heading("Good Leaver", level=3)

add_numbered(1, "Ein Good Leaver ist ein Gesellschafter, der aus einem der folgenden Gründe ausscheidet:")
add_bullet("Einvernehmliche Beendigung durch Beschluss aller Gesellschafter")
add_bullet("Dauerhafte Erkrankung oder Erwerbsunfähigkeit (ärztlich nachgewiesen)")
add_bullet("Tod")
add_bullet("Persönliche Gründe, die nicht unter Bad Leaver fallen, mit Ankündigungsfrist von 6 Monaten")

add_numbered(2, "Folgen für den Good Leaver:")
add_bullet("Bereits gevestete Anteile werden zum fairen Marktwert abgefunden")
add_bullet("Ungevestete Anteile fallen entschädigungslos an die Gesellschaft zurück")
add_bullet("Der faire Marktwert wird durch einen von der IHK benannten Wirtschaftsprüfer ermittelt, sofern sich die Gesellschafter nicht innerhalb von 30 Tagen auf einen Wert einigen")
add_bullet("Die Abfindung ist in drei gleichen Jahresraten zu zahlen, beginnend 6 Monate nach Ausscheiden")

doc.add_heading("Bad Leaver", level=3)

add_numbered(3, "Ein Bad Leaver ist ein Gesellschafter, der aus einem der folgenden Gründe ausscheidet:")
add_bullet("Fristlose Kündigung / Abberufung als Geschäftsführer aus wichtigem Grund")
add_bullet("Verstoß gegen das Wettbewerbsverbot (§ 12)")
add_bullet("Vorsätzliche Schädigung der Gesellschaft")
add_bullet("Rechtskräftige Verurteilung wegen einer Straftat im Zusammenhang mit der Gesellschaft")
add_bullet("Verletzung wesentlicher Pflichten aus diesem Vertrag trotz schriftlicher Abmahnung")

add_numbered(4, "Folgen für den Bad Leaver:")
add_bullet("Alle Anteile (gevestet und ungevestet) werden zum Nennwert (nicht Marktwert) eingezogen")
add_bullet("Der Nennwert beträgt den auf den Geschäftsanteil entfallenden Betrag des Stammkapitals (100,00 EUR)")
add_bullet("Die Zahlung ist sofort fällig")

doc.add_heading("Einziehung", level=3)

add_numbered(5, "Die Einziehung von Geschäftsanteilen ist mit Zustimmung des betroffenen Gesellschafters jederzeit zulässig, ohne seine Zustimmung nur aus den in diesem Vertrag genannten Gründen.")
add_numbered(6, "Die verbleibenden Gesellschafter sind berechtigt (nicht verpflichtet), die eingezogenen Anteile anteilig zu übernehmen.")

# ════════════════════════════════════════
# § 9
# ════════════════════════════════════════
doc.add_heading("§ 9 — Verfügung über Geschäftsanteile", level=2)

add_numbered(1, "Vorkaufsrecht: Bei beabsichtigter Übertragung von Geschäftsanteilen an Dritte haben die übrigen Gesellschafter ein Vorkaufsrecht zu denselben Bedingungen. Das Vorkaufsrecht ist innerhalb von 30 Tagen nach schriftlicher Mitteilung auszuüben.")
add_numbered(2, "Zustimmungspflicht: Jede Verfügung über Geschäftsanteile (Verkauf, Schenkung, Verpfändung) bedarf der Zustimmung aller übrigen Gesellschafter.")
add_numbered(3, "Tag-Along (Mitverkaufsrecht): Verkauft ein Gesellschafter seine Anteile an einen Dritten, haben die übrigen Gesellschafter das Recht, ihre Anteile zu denselben Konditionen mitzuverkaufen.")
add_numbered(4, "Drag-Along (Mitverkaufspflicht): Wollen Gesellschafter, die zusammen mindestens 75 % der Anteile halten, sämtliche Anteile an einen Dritten verkaufen, können sie die übrigen Gesellschafter verpflichten, ihre Anteile zu denselben Konditionen mitzuverkaufen. Der Kaufpreis muss mindestens dem fairen Marktwert (ermittelt wie in § 8 Abs. 2) entsprechen.")

# ════════════════════════════════════════
# § 10
# ════════════════════════════════════════
doc.add_heading("§ 10 — Geistiges Eigentum (Intellectual Property)", level=2)

add_numbered(1, "Sämtliche Arbeitsergebnisse, die von den Gesellschaftern oder Geschäftsführern im Rahmen ihrer Tätigkeit für die Gesellschaft erstellt werden, sind alleiniges Eigentum der Gesellschaft. Dies umfasst insbesondere:")
add_bullet("Quellcode, Software, Algorithmen")
add_bullet("Designs, Grafiken, UI/UX-Entwürfe")
add_bullet("Konzepte, Geschäftsmodelle, Strategiepapiere")
add_bullet("Dokumentationen, Texte, Marketingmaterialien")
add_bullet("Datenbanken und Datenstrukturen")
add_bullet("Domainnamen und Social-Media-Accounts der Gesellschaft")

add_numbered(2, "Vorbestehendes geistiges Eigentum (Pre-existing IP), das ein Gesellschafter vor Gründung der Gesellschaft erstellt hat und in die Gesellschaft einbringt, wird in Anlage 1 aufgelistet. Der jeweilige Gesellschafter räumt der Gesellschaft daran ein unwiderrufliches, übertragbares, weltweites Nutzungsrecht ein.")

add_numbered(3, "Jeder Gesellschafter versichert, dass er berechtigt ist, die von ihm eingebrachten Arbeitsergebnisse der Gesellschaft zu übertragen und dass keine Rechte Dritter entgegenstehen.")

add_numbered(4, "Jeder Gesellschafter ist verpflichtet, an allen erforderlichen Maßnahmen mitzuwirken, um der Gesellschaft die Rechte an den Arbeitsergebnissen zu sichern (z. B. Markenanmeldungen, Patente).")

# ════════════════════════════════════════
# § 11
# ════════════════════════════════════════
doc.add_heading("§ 11 — Geheimhaltung", level=2)

add_numbered(1, "Die Gesellschafter verpflichten sich, über alle vertraulichen Informationen der Gesellschaft Stillschweigen zu bewahren. Dies gilt insbesondere für:")
add_bullet("Geschäfts- und Betriebsgeheimnisse")
add_bullet("Kundendaten und Geschäftsbeziehungen")
add_bullet("Technische Informationen und Quellcode")
add_bullet("Finanzielle Informationen")
add_bullet("Strategische Pläne")

add_numbered(2, "Die Geheimhaltungspflicht gilt während der Gesellschafterstellung und 36 Monate nach dem Ausscheiden.")

add_numbered(3, "Ausgenommen sind Informationen, die:")
add_bullet("öffentlich bekannt sind (ohne Verschulden des Gesellschafters)")
add_bullet("von Dritten rechtmäßig erhalten wurden")
add_bullet("aufgrund gesetzlicher Verpflichtung offenzulegen sind")

# ════════════════════════════════════════
# § 12
# ════════════════════════════════════════
doc.add_heading("§ 12 — Wettbewerbsverbot", level=2)

add_numbered(1, "Während der Dauer der Gesellschafterstellung ist es jedem Gesellschafter untersagt, ohne schriftliche Zustimmung der übrigen Gesellschafter:")
add_bullet("ein konkurrierendes Unternehmen zu gründen, zu erwerben oder sich daran zu beteiligen")
add_bullet("für ein konkurrierendes Unternehmen tätig zu werden (angestellt oder freiberuflich)")
add_bullet("Kunden oder Mitarbeiter der Gesellschaft abzuwerben")

add_numbered(2, "Konkurrierend im Sinne dieses Vertrags ist jedes Unternehmen, das Software für Kundenbeziehungsmanagement (CRM) oder vergleichbare Unternehmenssoftware für kleine und mittlere Unternehmen im DACH-Raum entwickelt oder vertreibt.")

add_numbered(3, "Das Wettbewerbsverbot gilt für 12 Monate nach Ausscheiden aus der Gesellschaft. Für die Dauer des nachvertraglichen Wettbewerbsverbots zahlt die Gesellschaft eine Karenzentschädigung in Höhe von 50 % der letzten durchschnittlichen monatlichen Entnahmen/Vergütung des ausgeschiedenen Gesellschafters.")

add_numbered(4, "Ausnahme: Die Ausübung einer anderweitigen beruflichen Tätigkeit, die nicht im Wettbewerb zur Gesellschaft steht, ist jederzeit zulässig. Die Gesellschafter informieren sich gegenseitig über Nebentätigkeiten.")

# ════════════════════════════════════════
# § 13
# ════════════════════════════════════════
doc.add_heading("§ 13 — Gewinnverwendung und Rücklagen", level=2)

add_numbered(1, "Gesetzliche Rücklage (§ 5a GmbHG): Von dem um einen Verlustvortrag aus dem Vorjahr geminderten Jahresüberschuss ist ein Viertel (25 %) in eine gesetzliche Rücklage einzustellen, bis das Stammkapital den Betrag von 25.000,00 EUR erreicht hat.")
add_numbered(2, "Über die Verwendung des verbleibenden Gewinns beschließt die Gesellschafterversammlung mit einfacher Mehrheit.")
add_numbered(3, "Gewinnausschüttungen erfolgen im Verhältnis der Geschäftsanteile.")
add_numbered(4, "Die Gesellschafter streben an, in den ersten 24 Monaten nach Gründung die Gewinne vorrangig zu thesaurieren (im Unternehmen zu belassen), um das Wachstum zu finanzieren.")

# ════════════════════════════════════════
# § 14
# ════════════════════════════════════════
doc.add_heading("§ 14 — Vergütung der Geschäftsführer", level=2)

add_numbered(1, "Die Geschäftsführer üben ihre Tätigkeit zunächst unentgeltlich (ehrenamtlich) aus, solange die Gesellschaft keine ausreichenden Umsätze erzielt.")
add_numbered(2, "Sobald die Gesellschaft nachhaltige monatliche Umsätze von mindestens _____________ EUR [Betrag eintragen] erzielt, beschließt die Gesellschafterversammlung über eine angemessene Vergütung der Geschäftsführer.")
add_numbered(3, "Die Vergütung aller Geschäftsführer soll gleich hoch sein, sofern nicht einstimmig etwas anderes beschlossen wird.")
add_numbered(4, "Anpassungen der Vergütung bedürfen eines Gesellschafterbeschlusses mit einfacher Mehrheit.")

# ════════════════════════════════════════
# § 15
# ════════════════════════════════════════
doc.add_heading("§ 15 — Deadlock-Regelung", level=2)

add_numbered(1, "Können die Gesellschafter in einer wesentlichen Angelegenheit trotz zweimaliger Versammlung innerhalb von 30 Tagen keinen Beschluss fassen (Deadlock), gilt folgendes Verfahren:")

add_numbered(2, "Stufe 1 — Mediation (30 Tage): Die Gesellschafter benennen gemeinsam einen Mediator. Können sie sich nicht auf einen Mediator einigen, benennt die örtliche IHK einen Mediator. Die Kosten trägt die Gesellschaft.")

add_numbered(3, "Stufe 2 — Schiedsgutachten (30 Tage): Scheitert die Mediation, wird ein von der IHK benannter Schiedsgutachter mit der Entscheidung betraut. Sein Gutachten ist bindend.")

add_numbered(4, "Stufe 3 — Auflösungsrecht: Führen Stufe 1 und 2 nicht zu einer Lösung, hat jeder Gesellschafter das Recht, die Auflösung der Gesellschaft zu verlangen. Vor der Auflösung haben die übrigen Gesellschafter das Recht, die Anteile des auflösungswilligen Gesellschafters zum fairen Marktwert zu erwerben (Ankaufsrecht, Frist: 60 Tage).")

# ════════════════════════════════════════
# § 16
# ════════════════════════════════════════
doc.add_heading("§ 16 — Kündigung und Austritt", level=2)

add_numbered(1, "Jeder Gesellschafter kann seine Beteiligung mit einer Frist von 6 Monaten zum Ende des Geschäftsjahres kündigen.")
add_numbered(2, "Die Kündigung bedarf der Schriftform.")
add_numbered(3, "Im Falle der Kündigung gelten die Regelungen des § 8 (Good Leaver) entsprechend.")
add_numbered(4, "Die verbleibenden Gesellschafter sind berechtigt, die Gesellschaft fortzuführen. Die Geschäftsanteile des ausscheidenden Gesellschafters werden gemäß § 8 behandelt.")

# ════════════════════════════════════════
# § 17
# ════════════════════════════════════════
doc.add_heading("§ 17 — Umwandlung in GmbH", level=2)

add_numbered(1, "Die Gesellschafter streben an, die Gesellschaft in eine GmbH umzuwandeln, sobald die gesetzlichen Rücklagen zusammen mit gegebenenfalls zusätzlich aufzubringenden Mitteln den Betrag von 25.000,00 EUR Stammkapital erreichen.")
add_numbered(2, "Bei der Umwandlung bleiben die Beteiligungsverhältnisse, Vesting-Regelungen und alle übrigen Bestimmungen dieses Vertrags unverändert bestehen, soweit sie nicht im Widerspruch zur GmbH-Satzung stehen.")
add_numbered(3, "Die Kosten der Umwandlung trägt die Gesellschaft.")

# ════════════════════════════════════════
# § 18
# ════════════════════════════════════════
doc.add_heading("§ 18 — Auflösung und Liquidation", level=2)

add_numbered(1, "Die Gesellschaft wird aufgelöst durch:")
add_bullet("einstimmigen Beschluss der Gesellschafter")
add_bullet("gerichtliche Entscheidung")
add_bullet("Eröffnung des Insolvenzverfahrens")

add_numbered(2, "Im Falle der Auflösung sind die Geschäftsführer die Liquidatoren, sofern die Gesellschafterversammlung nicht andere Liquidatoren bestellt.")
add_numbered(3, "Nach Begleichung aller Verbindlichkeiten wird das verbleibende Vermögen im Verhältnis der Geschäftsanteile an die Gesellschafter verteilt.")

# ════════════════════════════════════════
# § 19
# ════════════════════════════════════════
doc.add_heading("§ 19 — Schlussbestimmungen", level=2)

add_numbered(1, "Änderungen und Ergänzungen dieses Vertrags bedürfen der notariellen Form.")
add_numbered(2, "Sollte eine Bestimmung dieses Vertrags unwirksam oder undurchführbar sein, wird die Wirksamkeit der übrigen Bestimmungen davon nicht berührt. An die Stelle der unwirksamen Bestimmung tritt eine wirksame Bestimmung, die dem wirtschaftlichen Zweck der unwirksamen Bestimmung am nächsten kommt (salvatorische Klausel).")
add_numbered(3, "Gerichtsstand für alle Streitigkeiten aus diesem Vertrag ist _______________________ [Stadt eintragen].")
add_numbered(4, "Es gilt das Recht der Bundesrepublik Deutschland.")

doc.add_page_break()

# ════════════════════════════════════════
# ANLAGE 1
# ════════════════════════════════════════
doc.add_heading("Anlage 1 — Vorbestehendes geistiges Eigentum", level=1)
doc.add_paragraph()

add_paragraph("Luke [Nachname]", bold=True)
add_table(
    ["Beschreibung", "Erstellungsdatum", "Lizenz an Gesellschaft"],
    [
        ["Backend-Codebase KMU Hub (vor Gründung)", "[Datum]", "Unwiderruflich, übertragbar, weltweit"],
        ["[weitere Einträge]", "", ""],
    ]
)

add_paragraph("Darien [Nachname]", bold=True)
add_table(
    ["Beschreibung", "Erstellungsdatum", "Lizenz an Gesellschaft"],
    [
        ["UI/UX-Designs, Frontend-Code (vor Gründung)", "[Datum]", "Unwiderruflich, übertragbar, weltweit"],
        ["[weitere Einträge]", "", ""],
    ]
)

add_paragraph("Nico [Nachname]", bold=True)
add_table(
    ["Beschreibung", "Erstellungsdatum", "Lizenz an Gesellschaft"],
    [
        ["QA-Dokumentation, Testprotokolle (vor Gründung)", "[Datum]", "Unwiderruflich, übertragbar, weltweit"],
        ["[weitere Einträge]", "", ""],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════
# ANLAGE 2
# ════════════════════════════════════════
doc.add_heading("Anlage 2 — Gesellschafterliste", level=1)
doc.add_paragraph()

add_table(
    ["Nr.", "Name", "Geburtsdatum", "Adresse", "Geschäftsanteil", "Nennbetrag"],
    [
        ["1", "Luke [Nachname]", "[Datum]", "[Adresse]", "33,3 %", "100,00 EUR"],
        ["2", "Darien [Nachname]", "[Datum]", "[Adresse]", "33,3 %", "100,00 EUR"],
        ["3", "Nico [Nachname]", "[Datum]", "[Adresse]", "33,3 %", "100,00 EUR"],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════
# UNTERSCHRIFTEN
# ════════════════════════════════════════
doc.add_heading("Unterschriften", level=1)
doc.add_paragraph()
doc.add_paragraph()

add_paragraph("Ort, Datum: _____________________________________________")
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

add_paragraph("_____________________________________________")
add_paragraph("Luke [Nachname]")
doc.add_paragraph()
doc.add_paragraph()

add_paragraph("_____________________________________________")
add_paragraph("Darien [Nachname]")
doc.add_paragraph()
doc.add_paragraph()

add_paragraph("_____________________________________________")
add_paragraph("Nico [Nachname]")

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Notariell beurkundet am _____________ durch Notar/in _____________")
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Urkundenrolle Nr. _____________")
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run.italic = True

# ── Save ──
output_path = r"C:\Users\darie\Documents\KMU Hub\KMU-Hub\.planning\legal\Gesellschaftervertrag_Zentria_UG.docx"
doc.save(output_path)
print(f"Saved to: {output_path}")
